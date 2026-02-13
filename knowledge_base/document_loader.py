"""
=============================================================================
DOCUMENT LOADER - Carregador de Documentos Multi-formato
=============================================================================

Este módulo contém funções para carregar diferentes tipos de documentos
que serão usados na base de conhecimento (RAG).

Tipos suportados:
- Arquivos de texto (.txt)
- Arquivos Markdown (.md)
- PDFs (.pdf) - requer pypdf
- CSVs (.csv) - usa pandas
- Word (.docx) - requer python-docx
- JSON (.json)

Conceito importante: CHUNKS
- Documentos grandes são divididos em pedaços menores (chunks)
- Cada chunk é indexado separadamente
- Isso permite busca mais precisa

=============================================================================
"""

from typing import List
from pathlib import Path
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
import io


# =============================================================================
# FORMATOS SUPORTADOS
# =============================================================================

SUPPORTED_FORMATS = {
    ".txt": "Texto simples",
    ".md": "Markdown",
    ".pdf": "PDF (requer pypdf)",
    ".csv": "CSV/Excel (requer pandas)",
    ".docx": "Word (requer python-docx)",
    ".json": "JSON",
}


def get_supported_formats() -> dict:
    """Retorna os formatos de arquivo suportados."""
    return SUPPORTED_FORMATS


# =============================================================================
# LOADERS POR TIPO DE ARQUIVO
# =============================================================================

def load_text_file(file_path: str = None, file_content: bytes = None,
                   filename: str = "unknown.txt", encoding: str = "utf-8") -> Document:
    """
    Carrega um arquivo de texto simples (.txt, .md).

    Args:
        file_path: Caminho para o arquivo (opcional se file_content for fornecido)
        file_content: Conteúdo do arquivo em bytes (para uploads)
        filename: Nome do arquivo (para metadados)
        encoding: Codificação do arquivo (padrão: utf-8)

    Returns:
        Um objeto Document do LangChain
    """
    if file_content is not None:
        content = file_content.decode(encoding)
    elif file_path is not None:
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"Arquivo não encontrado: {file_path}")
        content = path.read_text(encoding=encoding)
        filename = path.name
    else:
        raise ValueError("Forneça file_path ou file_content")

    metadata = {
        "source": file_path or filename,
        "filename": filename,
        "file_type": Path(filename).suffix
    }

    return Document(page_content=content, metadata=metadata)


def load_pdf_file(file_path: str = None, file_content: bytes = None,
                  filename: str = "unknown.pdf") -> List[Document]:
    """
    Carrega um arquivo PDF.

    Cada página do PDF se torna um Document separado.
    Requer: pip install pypdf

    Args:
        file_path: Caminho para o arquivo
        file_content: Conteúdo do arquivo em bytes
        filename: Nome do arquivo

    Returns:
        Lista de Documents (um por página)
    """
    try:
        from pypdf import PdfReader
    except ImportError:
        raise ImportError(
            "pypdf não está instalado. Execute: pip install pypdf"
        )

    if file_content is not None:
        pdf_file = io.BytesIO(file_content)
    elif file_path is not None:
        pdf_file = file_path
        filename = Path(file_path).name
    else:
        raise ValueError("Forneça file_path ou file_content")

    reader = PdfReader(pdf_file)
    documents = []

    for page_num, page in enumerate(reader.pages, 1):
        text = page.extract_text()
        if text.strip():  # Ignora páginas vazias
            doc = Document(
                page_content=text,
                metadata={
                    "source": file_path or filename,
                    "filename": filename,
                    "file_type": ".pdf",
                    "page": page_num,
                    "total_pages": len(reader.pages)
                }
            )
            documents.append(doc)

    return documents


def load_csv_file(file_path: str = None, file_content: bytes = None,
                  filename: str = "unknown.csv") -> List[Document]:
    """
    Carrega um arquivo CSV.

    Cada linha do CSV se torna um Document separado.
    Requer: pip install pandas

    Args:
        file_path: Caminho para o arquivo
        file_content: Conteúdo do arquivo em bytes
        filename: Nome do arquivo

    Returns:
        Lista de Documents (um por linha)
    """
    try:
        import pandas as pd
    except ImportError:
        raise ImportError(
            "pandas não está instalado. Execute: pip install pandas"
        )

    if file_content is not None:
        csv_file = io.BytesIO(file_content)
    elif file_path is not None:
        csv_file = file_path
        filename = Path(file_path).name
    else:
        raise ValueError("Forneça file_path ou file_content")

    df = pd.read_csv(csv_file)
    documents = []

    # Opção 1: Cada linha vira um documento
    for idx, row in df.iterrows():
        # Formata a linha como texto legível
        content = "\n".join([f"{col}: {val}" for col, val in row.items()])
        doc = Document(
            page_content=content,
            metadata={
                "source": file_path or filename,
                "filename": filename,
                "file_type": ".csv",
                "row": idx + 1,
                "total_rows": len(df)
            }
        )
        documents.append(doc)

    return documents


def load_docx_file(file_path: str = None, file_content: bytes = None,
                   filename: str = "unknown.docx") -> Document:
    """
    Carrega um arquivo Word (.docx).

    Requer: pip install python-docx

    Args:
        file_path: Caminho para o arquivo
        file_content: Conteúdo do arquivo em bytes
        filename: Nome do arquivo

    Returns:
        Um Document com todo o conteúdo do arquivo
    """
    try:
        from docx import Document as DocxDocument
    except ImportError:
        raise ImportError(
            "python-docx não está instalado. Execute: pip install python-docx"
        )

    if file_content is not None:
        docx_file = io.BytesIO(file_content)
    elif file_path is not None:
        docx_file = file_path
        filename = Path(file_path).name
    else:
        raise ValueError("Forneça file_path ou file_content")

    doc = DocxDocument(docx_file)

    # Extrai texto de todos os parágrafos
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    content = "\n\n".join(paragraphs)

    return Document(
        page_content=content,
        metadata={
            "source": file_path or filename,
            "filename": filename,
            "file_type": ".docx",
            "paragraphs": len(paragraphs)
        }
    )


def load_json_file(file_path: str = None, file_content: bytes = None,
                   filename: str = "unknown.json",
                   text_field: str = None) -> List[Document]:
    """
    Carrega um arquivo JSON.

    Se o JSON for uma lista, cada item vira um Document.
    Se for um objeto, todo o conteúdo vira um Document.

    Args:
        file_path: Caminho para o arquivo
        file_content: Conteúdo do arquivo em bytes
        filename: Nome do arquivo
        text_field: Campo específico para usar como conteúdo (opcional)

    Returns:
        Lista de Documents
    """
    import json

    if file_content is not None:
        data = json.loads(file_content.decode("utf-8"))
    elif file_path is not None:
        with open(file_path, "r", encoding="utf-8") as f:
            data = json.load(f)
        filename = Path(file_path).name
    else:
        raise ValueError("Forneça file_path ou file_content")

    documents = []

    def item_to_text(item):
        """Converte um item do JSON em texto."""
        if isinstance(item, dict):
            if text_field and text_field in item:
                return str(item[text_field])
            return "\n".join([f"{k}: {v}" for k, v in item.items()])
        return str(item)

    if isinstance(data, list):
        for idx, item in enumerate(data):
            doc = Document(
                page_content=item_to_text(item),
                metadata={
                    "source": file_path or filename,
                    "filename": filename,
                    "file_type": ".json",
                    "index": idx
                }
            )
            documents.append(doc)
    else:
        doc = Document(
            page_content=item_to_text(data),
            metadata={
                "source": file_path or filename,
                "filename": filename,
                "file_type": ".json"
            }
        )
        documents.append(doc)

    return documents


# =============================================================================
# FUNÇÃO UNIVERSAL DE CARREGAMENTO
# =============================================================================

def load_document(file_path: str = None, file_content: bytes = None,
                  filename: str = None) -> List[Document]:
    """
    Carrega um documento automaticamente baseado na extensão.

    Esta é a função principal - detecta o tipo de arquivo e usa
    o loader apropriado.

    Args:
        file_path: Caminho para o arquivo
        file_content: Conteúdo do arquivo em bytes (para uploads)
        filename: Nome do arquivo (necessário se usar file_content)

    Returns:
        Lista de Documents

    Example:
        >>> # Carregar de arquivo
        >>> docs = load_document(file_path="documento.pdf")

        >>> # Carregar de upload (Streamlit)
        >>> docs = load_document(file_content=uploaded_file.read(),
        ...                      filename=uploaded_file.name)
    """
    # Determina o nome do arquivo
    if filename is None and file_path is not None:
        filename = Path(file_path).name
    elif filename is None:
        raise ValueError("Forneça filename quando usar file_content")

    # Detecta extensão
    extension = Path(filename).suffix.lower()

    # Seleciona o loader apropriado
    if extension in [".txt", ".md"]:
        doc = load_text_file(file_path, file_content, filename)
        return [doc]

    elif extension == ".pdf":
        return load_pdf_file(file_path, file_content, filename)

    elif extension == ".csv":
        return load_csv_file(file_path, file_content, filename)

    elif extension == ".docx":
        doc = load_docx_file(file_path, file_content, filename)
        return [doc]

    elif extension == ".json":
        return load_json_file(file_path, file_content, filename)

    else:
        raise ValueError(
            f"Formato não suportado: {extension}\n"
            f"Formatos aceitos: {', '.join(SUPPORTED_FORMATS.keys())}"
        )


# =============================================================================
# FUNÇÕES DE DIRETÓRIO
# =============================================================================

def load_documents_from_directory(
    directory_path: str,
    file_extensions: List[str] = None,
    recursive: bool = True
) -> List[Document]:
    """
    Carrega todos os documentos de um diretório.

    Args:
        directory_path: Caminho para o diretório
        file_extensions: Lista de extensões a processar (None = todas suportadas)
        recursive: Se True, processa subdiretórios também

    Returns:
        Lista de Documents carregados
    """
    directory = Path(directory_path)

    if not directory.exists():
        raise FileNotFoundError(f"Diretório não encontrado: {directory_path}")

    if file_extensions is None:
        file_extensions = list(SUPPORTED_FORMATS.keys())

    documents = []
    pattern = "**/*" if recursive else "*"

    for file_path in directory.glob(pattern):
        if file_path.is_file() and file_path.suffix.lower() in file_extensions:
            try:
                docs = load_document(file_path=str(file_path))
                documents.extend(docs)
                print(f"✅ Carregado: {file_path.name}")
            except Exception as e:
                print(f"❌ Erro ao carregar {file_path.name}: {e}")

    return documents


# =============================================================================
# FUNÇÃO DE SPLIT (CHUNKING)
# =============================================================================

def split_documents(
    documents: List[Document],
    chunk_size: int = 1000,
    chunk_overlap: int = 200
) -> List[Document]:
    """
    Divide documentos em chunks menores.

    Por que dividir?
    - LLMs têm limite de tokens
    - Chunks menores = busca mais precisa
    - Overlap evita perda de contexto nas bordas

    Args:
        documents: Lista de documentos a dividir
        chunk_size: Tamanho máximo de cada chunk (em caracteres)
        chunk_overlap: Sobreposição entre chunks

    Returns:
        Lista de Documents divididos
    """
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len,
        separators=["\n\n", "\n", ". ", " ", ""]
    )

    chunks = text_splitter.split_documents(documents)
    print(f"📄 {len(documents)} documentos divididos em {len(chunks)} chunks")

    return chunks


# =============================================================================
# EXEMPLO DE USO
# =============================================================================
if __name__ == "__main__":
    print("🚀 Testando Document Loader Multi-formato")
    print("=" * 50)
    print(f"\nFormatos suportados: {list(SUPPORTED_FORMATS.keys())}")

    # Teste com arquivo de texto
    test_content = """
    # Documento de Teste
    
    Este é um documento de teste para o RAG.
    Contém informações importantes sobre o sistema.
    """

    test_file = Path("test_doc.txt")
    test_file.write_text(test_content)

    docs = load_document(file_path="test_doc.txt")
    print(f"\n✅ Carregado: {len(docs)} documento(s)")
    print(f"Conteúdo: {docs[0].page_content[:100]}...")

    test_file.unlink()



